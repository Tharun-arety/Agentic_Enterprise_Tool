from __future__ import annotations
import asyncio,io,re
from datetime import datetime,timezone
from pathlib import Path
from sqlalchemy import delete,func,select
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.embeddings import get_embedding_provider
from app.core.filestore import content_key,get_file_store
from app.domains.knowledge.models import DocumentChunk,SourceDocument,SourceDocumentVersion
from app.tools.registry import ToolError

SUPPORTED={".pdf",".docx",".xlsx",".html",".htm",".md",".markdown",".txt"}
async def ingest(session:AsyncSession,*,document_key:str,title:str,filename:str,content_type:str,data:bytes,acl_labels:list[str],revision:str|None,user_id=None):
    if not data: raise ToolError("Cannot ingest an empty document.")
    suffix=Path(filename).suffix.lower()
    if suffix not in SUPPORTED: raise ToolError(f"Unsupported document type {suffix!r}.")
    digest,key=content_key(data,filename); doc=(await session.execute(select(SourceDocument).where(SourceDocument.document_key==document_key))).scalar_one_or_none()
    if doc is None: doc=SourceDocument(document_key=document_key,title=title,acl_labels=acl_labels or ["internal"]); session.add(doc); await session.flush()
    duplicate=(await session.execute(select(SourceDocumentVersion).where(SourceDocumentVersion.source_document_id==doc.id,SourceDocumentVersion.content_hash==digest))).scalar_one_or_none()
    if duplicate: return duplicate,False
    version=int(await session.scalar(select(func.count()).select_from(SourceDocumentVersion).where(SourceDocumentVersion.source_document_id==doc.id)) or 0)+1
    await get_file_store().put(key,data,content_type); row=SourceDocumentVersion(source_document_id=doc.id,version=version,revision=revision,filename=filename,content_type=content_type,content_hash=digest,storage_key=key,uploaded_by=user_id); session.add(row); await session.flush()
    try:
        sections=await asyncio.to_thread(_extract,suffix,data)
        texts=[]; metas=[]
        for text,meta in sections:
            for chunk in _chunks(text): texts.append(chunk); metas.append(meta)
        vectors=await get_embedding_provider().embed(texts)
        for i,(text,meta,vector) in enumerate(zip(texts,metas,vectors,strict=True)): session.add(DocumentChunk(document_version_id=row.id,ordinal=i,text_content=text,embedding=vector,page=meta.get("page"),sheet=meta.get("sheet"),heading=meta.get("heading"),metadata_json=meta))
        row.extraction_status="indexed"; row.extracted_at=datetime.now(timezone.utc)
    except Exception as exc:
        row.extraction_status="failed"; row.extraction_error=f"{type(exc).__name__}: {exc}"[:2000]
    await session.commit(); return row,True

async def reindex(session:AsyncSession,version_id):
    row=await session.get(SourceDocumentVersion,version_id)
    if not row: raise ToolError("Document version does not exist.")
    data=await get_file_store().get(row.storage_key); await session.execute(delete(DocumentChunk).where(DocumentChunk.document_version_id==row.id)); row.extraction_status="pending"; await session.flush()
    sections=await asyncio.to_thread(_extract,Path(row.filename).suffix.lower(),data); texts=[]; metas=[]
    for text,meta in sections:
        for chunk in _chunks(text): texts.append(chunk); metas.append(meta)
    for i,(text,meta,vector) in enumerate(zip(texts,metas,await get_embedding_provider().embed(texts),strict=True)): session.add(DocumentChunk(document_version_id=row.id,ordinal=i,text_content=text,embedding=vector,page=meta.get("page"),sheet=meta.get("sheet"),heading=meta.get("heading"),metadata_json=meta))
    row.extraction_status="indexed"; row.extracted_at=datetime.now(timezone.utc); row.extraction_error=None; await session.commit(); return row

def _chunks(text,size=1200,overlap=160):
    clean=re.sub(r"\s+"," ",text).strip(); out=[]; start=0
    while start<len(clean):
        end=min(len(clean),start+size); out.append(clean[start:end]); start=end-overlap if end<len(clean) else end
    return out
def _extract(suffix,data):
    if suffix==".pdf":
        from pypdf import PdfReader
        reader=PdfReader(io.BytesIO(data)); rows=[(p.extract_text() or "",{"page":i+1}) for i,p in enumerate(reader.pages)]
        if not any(t.strip() for t,_ in rows): return _ocr_pdf(data)
        return rows
    if suffix==".docx":
        from docx import Document
        doc=Document(io.BytesIO(data)); rows=[]; heading=None; buf=[]
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith("Heading"):
                if buf: rows.append(("\n".join(buf),{"heading":heading})); buf=[]
                heading=p.text
            else: buf.append(p.text)
        if buf: rows.append(("\n".join(buf),{"heading":heading}))
        return rows
    if suffix==".xlsx":
        from openpyxl import load_workbook
        wb=load_workbook(io.BytesIO(data),read_only=True,data_only=True); return [("\n".join("\t".join(str(c) if c is not None else "" for c in row) for row in ws.iter_rows(values_only=True)),{"sheet":ws.title}) for ws in wb.worksheets]
    text=data.decode("utf-8",errors="replace")
    if suffix in {".html",".htm"}:
        from bs4 import BeautifulSoup
        soup=BeautifulSoup(text,"html.parser"); return [(soup.get_text(" "),{"heading":soup.title.string if soup.title else None})]
    if suffix in {".md",".markdown"}:
        rows=[]; heading=None; buf=[]
        for line in text.splitlines():
            if line.startswith("#"):
                if buf: rows.append(("\n".join(buf),{"heading":heading})); buf=[]
                heading=line.lstrip("# ")
            else: buf.append(line)
        if buf: rows.append(("\n".join(buf),{"heading":heading}))
        return rows
    return [(text,{})]
def _ocr_pdf(data):
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        return [(pytesseract.image_to_string(image),{"page":i+1,"ocr":True}) for i,image in enumerate(convert_from_bytes(data))]
    except Exception as exc: raise RuntimeError("PDF contains no extractable text and OCR fallback requires Tesseract plus Poppler") from exc
